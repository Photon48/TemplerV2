import streamlit as st

from dotenv import load_dotenv
from langchain.document_loaders import DirectoryLoader
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chat_models import ChatOpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain import OpenAI
from langchain.chains import RetrievalQA
import magic
from io import BytesIO
import tempfile
import os
from docx import Document
import zipfile
import nltk
import pypdfium2 as pdfium
import matplotlib.pyplot as plt
from PIL import Image
from pytesseract import image_to_string 
from langchain.chains.summarize import load_summarize_chain


# Prompt templates for dynamic values
from langchain.prompts.chat import (
    ChatPromptTemplate,
    SystemMessagePromptTemplate,
    AIMessagePromptTemplate, # I included this one so you know you'll have it but we won't be using it
    HumanMessagePromptTemplate
)

# To create our chat messages
from langchain.schema import (
    AIMessage,
    HumanMessage,
    SystemMessage
)

openai_api_key = st.secrets["OPENAI_API_KEY"]

nltk.download('averaged_perceptron_tagger')



def convert_pdf_to_images(file_path, scale=300/72):
    
    pdf_file = pdfium.PdfDocument(file_path)  
    page_indices = [i for i in range(len(pdf_file))]
    
    renderer = pdf_file.render(
        pdfium.PdfBitmap.to_pil,
        page_indices = page_indices, 
        scale = scale,
    )
    
    list_final_images = [] 
    
    for i, image in zip(page_indices, renderer):
        
        image_byte_array = BytesIO()
        image.save(image_byte_array, format='jpeg', optimize=True)
        image_byte_array = image_byte_array.getvalue()
        list_final_images.append(dict({i:image_byte_array}))
    
    return list_final_images

def extract_text_with_pytesseract(list_dict_final_images):
    
    image_list = [list(data.values())[0] for data in list_dict_final_images]
    image_content = []
    
    for index, image_bytes in enumerate(image_list):
        
        image = Image.open(BytesIO(image_bytes))
        raw_text = str(image_to_string(image))
        image_content.append(raw_text)
    
    return "\n".join(image_content)

# def convert_docx_to_txt(file_bytes):
#     # Create a Document object from the bytes
#     doc = Document(BytesIO(file_bytes))

#     # Extract text from each paragraph and join it into one string
#     text = ' '.join(paragraph.text for paragraph in doc.paragraphs)

#     return text


# def convert_txt_to_txt(path): ##HAVING PROBLEMS WITH TXT CONVERSION FOR SOME GODDAMN REASON!
#     with open(path, 'r') as file:
#         text = file.read()

#     return text

def convert_uploaded_files_to_text(uploaded_files):
    text_files = {}

    for uploaded_file in uploaded_files:
        # Get the file extension
        _, extension = os.path.splitext(uploaded_file.name)

        # Read the file bytes
        file_bytes = uploaded_file.getvalue()

        if extension == '.pdf':
            # Create a temporary file with the same extension
            with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as temp:
                # Write the bytes to the temporary file
                temp.write(file_bytes)
                # Pass the path of the temporary file to your function
                pdf_to_image = convert_pdf_to_images(temp.name)
            # Remember to remove the temporary file after use
            os.remove(temp.name)
            
            text_files[uploaded_file.name] = extract_text_with_pytesseract(pdf_to_image)
        else:
            print(f"Unsupported file format: {uploaded_file.name}")

    return text_files

# Set the title of the Streamlit application
st.title("RAG-based Question Answering App")

    # Create a slider to let the user select how many blocks they want to create
    # The slider ranges from 1 to 30, with a default value of 1
num_blocks = st.slider("Number of blocks", 1, 30, 1)

with st.form("file_upload_form"):
    # Initialize an empty list to store the blocks
    blocks = []

    # For each block that the user wants to create
    for i in range(num_blocks):
        # Display a subheader that shows which block the user is creating
        st.subheader(f'Block {i+1}')

        # Create a dropdown menu where the user can select the type of the block
        block_type = st.selectbox(f'Choose the type of block for Block {i+1}:', ('Title', 'Question', 'Text'), key=f'block_type{i}')

        # Create a single text input field for each block
        content = st.text_input('Enter the content', key=f'content{i}')

        # Add the block to the list of blocks
        blocks.append({'type': block_type, 'content': content})

    uploaded_files = st.file_uploader("Choose your source files", accept_multiple_files=True)

    # Create a submit button
    if st.form_submit_button('Submit'):
        with st.spinner('Processing...'):
            if uploaded_files:
                text_files = convert_uploaded_files_to_text(uploaded_files)
                transcript = ' '.join(text_files.values())
                # Load up your text splitter
                text_splitter = RecursiveCharacterTextSplitter(separators=["\n\n", "\n", " "], chunk_size=2000, chunk_overlap=0)
                # I'm only doing the first 23250 characters. This to save on costs. When you're doing your exercise you can remove this to let all the data through
                #transcript_subsection_characters = 23250
                docs = text_splitter.create_documents([transcript])
                print(f"You have {len(docs)} docs.")
                texts = text_splitter.create_documents([transcript])
                embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
                # Get your docsearch ready
                docsearch = FAISS.from_documents(texts, embeddings)
                # Load up your LLM
                llm = OpenAI(openai_api_key=openai_api_key)
                qa = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=docsearch.as_retriever())
                # Initialize an empty list to store the answers
                answers = []
                # Initialize a dictionary to store the blocks by type
                blocks_by_type = {'Title': [], 'Question': [], 'Text': []}
                # Process each block individually
                # Initialize an empty list to store the processed blocks
                processed_blocks = []
                # Process each block individually
                for block in blocks:
                    # If the block is a question, run it through the QA model and get the answer
                    if block['type'] == 'Question':
                        query = str(block['content'])
                        answer = qa.run(query)
                        processed_blocks.append({'type': 'Answer', 'content': answer})
                    else:
                        processed_blocks.append(block)

                # Display the processed blocks
                for block in processed_blocks:
                    if block['type'] == 'Title':
                        st.header(block['content'])
                    elif block['type'] == 'Answer':
                        st.write(block['content'])  # This will display the answer
                    else:  # block['type'] == 'Text'
                        st.write(block['content']) 